# -*- coding: utf-8 -*-

from dotenv import load_dotenv
import os
from openai import OpenAI
import httpx
from docx import Document  # 用于读取 Word 文件作为 prompt
import json
import re

# 加载环境变量
load_dotenv()
api_key = os.getenv("API_KEY")


# 加载系统提示（Prompt）文本
def load_prompt_from_docx(path: str) -> str:
    try:
        doc = Document(path)
        return "\n".join([p.text.strip() for p in doc.paragraphs if p.text.strip()])
    except Exception as e:
        print(f"[Prompt Load Error]: {e}")
        return "You are an AI assistant. Please respond politely."


system_prompt = load_prompt_from_docx("./key/deepseek_prompt.docx")

# 初始化客户端
client = OpenAI(
    api_key=api_key,
    base_url="https://api.deepseek.com",
    http_client=httpx.Client(timeout=60.0)
)


# 加载品牌映射表
def load_brand_mapping():
    try:
        with open('./backend/data/brand_mapping.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"[Brand Mapping Load Error]: {e}")
        return {}


# 加载商品类型映射表
def load_category_mapping():
    try:
        with open('./backend/data/category_mapping.json', 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"[Category Mapping Load Error]: {e}")
        return {}


# 全局加载品牌映射
brand_mapping = load_brand_mapping()
category_mapping = load_category_mapping()


# 清理系列名称（去除"系列"、括号内容等）
def clean_series_name(series_name: str) -> str:
    # 去除Markdown格式（如 **text** 或 *text*）
    series_name = re.sub(r'\*+([^*]+)\*+', r'\1', series_name)
    # 去除括号及其内容
    series_name = re.sub(r'[（(][^）)]*[）)]', '', series_name)
    # 去除"系列"二字
    series_name = series_name.replace('系列', '')
    # 去除多余空格
    return series_name.strip()


# 检查字符串是否包含中文
def contains_chinese(text: str) -> bool:
    return bool(re.search(r'[\u4e00-\u9fff]', text))


# 翻译商品类型
def translate_category(category: str) -> str:
    """翻译商品类型为英文"""
    # 先查找映射表
    if category.lower() in category_mapping:
        return category_mapping[category.lower()]

    # 如果不包含中文，直接返回
    if not contains_chinese(category):
        return category

    # 调用API翻译
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system",
                 "content": "You are a translator. Translate product categories to English. Only return the English translation, nothing else."},
                {"role": "user", "content": f"Translate this product category to English: {category}"}
            ],
            stream=False
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"[Category Translation Error]: {e}")
        return category  # 翻译失败时返回原文


# 生成搜索关键词
def generate_search_keyword(brand_name: str, series_name: str, category: str = "", language: str = "zh") -> str:
    """
    生成适合俄罗斯电商平台搜索的英文关键词
    """
    # 清理系列名称
    clean_series = clean_series_name(series_name)

    # 1. 检查品牌名是否在映射表中
    english_brand = brand_mapping.get(brand_name, None)

    # 2. 如果语言不是中文，需要确保系列名也是英文
    if language != "zh":
        # 尝试从系列名中提取英文部分
        english_match = re.search(r'[（(]([A-Za-z\s\-]+)[）)]', clean_series)
        if english_match:
            clean_series = english_match[1].strip()
        else:
            # 检查是否已经是纯英文
            if not contains_chinese(clean_series):
                # 已经是英文，直接使用
                pass
            else:
                # 需要翻译
                query = f"请将'{clean_series}'翻译成英文，只返回英文翻译，不要其他内容"
                try:
                    response = client.chat.completions.create(
                        model="deepseek-chat",
                        messages=[
                            {"role": "system", "content": "你是一个翻译助手，只返回英文翻译"},
                            {"role": "user", "content": query}
                        ],
                        stream=False
                    )
                    clean_series = response.choices[0].message.content.strip()
                except Exception as e:
                    print(f"[Series Translation Error]: {e}")

    # 3. 如果品牌名不在映射表或系列名包含中文，调用DeepSeek API
    if not english_brand or contains_chinese(clean_series):
        # 构建请求
        query = f"请将'{brand_name} {clean_series}'转换为适合在俄罗斯电商搜索的英文关键词，只返回'英文品牌 英文系列'格式，不要其他内容"

        try:
            response = client.chat.completions.create(
                model="deepseek-chat",
                messages=[
                    {"role": "system", "content": "你是一个翻译助手，只返回英文品牌名和系列名，格式：Brand Series"},
                    {"role": "user", "content": query}
                ],
                stream=False
            )
            result = response.choices[0].message.content.strip()
            # 解析返回的品牌和系列
            parts = result.split()
            if len(parts) >= 2:
                english_brand = parts[0]
                clean_series = ' '.join(parts[1:])
            else:
                english_brand = result
                clean_series = ""
        except Exception as e:
            print(f"[DeepSeek Translation Error]: {e}")
            # 如果API调用失败，使用原始值
            if not english_brand:
                english_brand = brand_name

    # 4. 检查品牌名和系列名是否相同（忽略大小写）
    if english_brand and clean_series and english_brand.lower() == clean_series.lower():
        clean_series = ""  # 如果相同，忽略系列名

    # 5. 处理商品类型
    english_category = ""
    if category:
        english_category = translate_category(category)

    # 6. 组合最终的搜索关键词
    keyword_parts = [english_brand]
    if clean_series:
        keyword_parts.append(clean_series)
    if english_category:
        keyword_parts.append(english_category)

    return ' '.join(keyword_parts).strip()


# 语言指令映射
LANGUAGE_INSTRUCTIONS = {
    "zh": "\n\n请用中文回答。",
    "en": "\n\nPlease respond in English.",
    "ru": "\n\nПожалуйста, отвечайте на русском языке."
}


# 请求函数
def call_deepseek_api(query: str, language: str = "zh") -> str:
    print(f"[DeepSeek API] search request received: {query}, language: {language}")
    try:
        # 添加语言指令
        lang_instruction = LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["zh"])

        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": system_prompt + lang_instruction},
                {"role": "user", "content": query}
            ],
            stream=False
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print("[DeepSeek API Error]:", e)
        return "❌ 调用 DeepSeek API 时出错，请稍后再试。"


# 流式请求函数
def call_deepseek_api_stream(query: str, language: str = "zh"):
    """流式调用 DeepSeek API"""
    print(f"[DeepSeek API] stream request received: {query}, language: {language}")
    try:
        # 添加语言指令
        lang_instruction = LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["zh"])

        stream = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": system_prompt + lang_instruction},
                {"role": "user", "content": query}
            ],
            stream=True  # 启用流式输出
        )

        for chunk in stream:
            if chunk.choices[0].delta.content is not None:
                yield chunk.choices[0].delta.content

    except Exception as e:
        print("[DeepSeek API Stream Error]:", e)
        yield f"❌ 调用 DeepSeek API 时出错：{str(e)}"


def call_deepseek_brand_prompt(brand_name: str, language: str = "zh") -> str:
    from docx import Document
    try:
        document = Document("key/deepseek_brand_prompt.docx")
        prompt_text = "\n".join([para.text for para in document.paragraphs])
    except:
        # 如果文档不存在，使用内置的提示词
        prompt_text = """你将被用于回答某个具体品牌的详细特征。请根据输入的品牌名，输出品牌的以下内容：
- 品牌简介
- 优势与劣势
- 同类竞品对比（以表格呈现）
- 总结建议
- 该品牌的官方网站

在回答的最后，请按以下格式列出该品牌最具代表性的5个产品系列（如果不足5个则列出全部）：
◆◆◆ 产品系列1
◆◆◆ 产品系列2
◆◆◆ 产品系列3
◆◆◆ 产品系列4
◆◆◆ 产品系列5"""

    # 替换提示词中的标记符号
    prompt_text = prompt_text.replace("🔸", "◆◆◆")

    # 添加语言指令
    lang_instruction = LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["zh"])

    # 如果是非中文语言，添加额外说明
    if language != "zh":
        extra_instruction = "\n\nIMPORTANT: List product series names in English. For example: 'Fresh Series', 'Smart Series', etc."
        full_prompt = prompt_text + f"\n\n品牌名称：{brand_name}" + extra_instruction + lang_instruction
    else:
        full_prompt = prompt_text + f"\n\n品牌名称：{brand_name}" + lang_instruction

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "user", "content": full_prompt}
        ]
    )
    return response.choices[0].message.content


def call_deepseek_brand_prompt_stream(brand_name: str, language: str = "zh"):
    """流式调用品牌详情API"""
    from docx import Document
    try:
        document = Document("key/deepseek_brand_prompt.docx")
        prompt_text = "\n".join([para.text for para in document.paragraphs])
    except:
        # 如果文档不存在，使用内置的提示词
        prompt_text = """你将被用于回答某个具体品牌的详细特征。请根据输入的品牌名，输出品牌的以下内容：
- 品牌简介
- 优势与劣势
- 同类竞品对比（以表格呈现）
- 总结建议
- 该品牌的官方网站

在回答的最后，请按以下格式列出该品牌最具代表性的5个产品系列（如果不足5个则列出全部）：
◆◆◆ 产品系列1
◆◆◆ 产品系列2
◆◆◆ 产品系列3
◆◆◆ 产品系列4
◆◆◆ 产品系列5"""

    # 替换提示词中的标记符号
    prompt_text = prompt_text.replace("🔸", "◆◆◆")

    # 添加语言指令
    lang_instruction = LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["zh"])

    # 如果是非中文语言，添加额外说明
    if language != "zh":
        extra_instruction = "\n\nIMPORTANT: List product series names in English. For example: 'Fresh Series', 'Smart Series', etc."
        full_prompt = prompt_text + f"\n\n品牌名称：{brand_name}" + extra_instruction + lang_instruction
    else:
        full_prompt = prompt_text + f"\n\n品牌名称：{brand_name}" + lang_instruction

    stream = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "user", "content": full_prompt}
        ],
        stream=True
    )

    for chunk in stream:
        if chunk.choices[0].delta.content is not None:
            yield chunk.choices[0].delta.content


def call_deepseek_brand_with_context_prompt(brand_name: str, context: str, language: str = "zh") -> str:
    """根据商品类型上下文生成品牌详情"""
    # 添加语言指令
    lang_instruction = LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["zh"])

    try:
        # 尝试加载专门的上下文提示词文档
        from docx import Document
        document = Document("key/deepseek_brand_context_prompt.docx")
        prompt_text = "\n".join([para.text for para in document.paragraphs])
        # 替换提示词中的标记符号
        prompt_text = prompt_text.replace("🔸", "◆◆◆")

        # 如果是非中文语言，添加额外说明
        if language != "zh":
            extra_instruction = "\n\nIMPORTANT: List product series names in English, not in Chinese. For example: 'Fresh Series', 'Smart Series', etc."
            full_prompt = prompt_text + f"\n\n品牌名称：{brand_name}\n商品类型：{context}" + extra_instruction + lang_instruction
        else:
            full_prompt = prompt_text + f"\n\n品牌名称：{brand_name}\n商品类型：{context}" + lang_instruction
    except:
        # 如果文档不存在，使用内置的提示词
        base_prompt = f"""
你将被用于回答某个具体品牌在特定商品类型下的详细信息。

品牌名称：{brand_name}
商品类型：{context}

请输出该品牌在{context}领域的以下内容：

- 品牌在{context}领域的简介
- 该品牌{context}产品的优势与劣势
- 与其他品牌{context}产品的对比（以表格呈现）
- 总结建议
- 该品牌的官方网站

在回答的最后，请按以下格式列出该品牌在{context}领域最具代表性的5个产品系列（如果不足5个则列出全部）：

◆◆◆ 产品系列1（仅限{context}相关）
◆◆◆ 产品系列2（仅限{context}相关）
◆◆◆ 产品系列3（仅限{context}相关）
◆◆◆ 产品系列4（仅限{context}相关）
◆◆◆ 产品系列5（仅限{context}相关）

重要提示：
1. 只列出与{context}相关的产品系列
2. 如果该品牌没有{context}相关产品，请明确说明
3. 不要列出其他类型的产品系列（如手机、电视等与{context}无关的产品）
4. 产品系列名称前必须使用"◆◆◆"标记，不要使用其他符号
"""
        if language != "zh":
            base_prompt += "\n5. Product series names must be in English, not Chinese"

        full_prompt = base_prompt.replace("请用中文回答。", "").strip() + lang_instruction

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "user", "content": full_prompt}
        ]
    )
    return response.choices[0].message.content


def call_deepseek_brand_with_context_prompt_stream(brand_name: str, context: str, language: str = "zh"):
    """流式调用带上下文的品牌详情API"""
    # 添加语言指令
    lang_instruction = LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["zh"])

    try:
        # 尝试加载专门的上下文提示词文档
        from docx import Document
        document = Document("key/deepseek_brand_context_prompt.docx")
        prompt_text = "\n".join([para.text for para in document.paragraphs])
        # 替换提示词中的标记符号
        prompt_text = prompt_text.replace("🔸", "◆◆◆")

        # 如果是非中文语言，添加额外说明
        if language != "zh":
            extra_instruction = "\n\nIMPORTANT: List product series names in English, not in Chinese. For example: 'Fresh Series', 'Smart Series', etc."
            full_prompt = prompt_text + f"\n\n品牌名称：{brand_name}\n商品类型：{context}" + extra_instruction + lang_instruction
        else:
            full_prompt = prompt_text + f"\n\n品牌名称：{brand_name}\n商品类型：{context}" + lang_instruction
    except:
        # 如果文档不存在，使用内置的提示词
        base_prompt = f"""
你将被用于回答某个具体品牌在特定商品类型下的详细信息。

品牌名称：{brand_name}
商品类型：{context}

请输出该品牌在{context}领域的以下内容：

- 品牌在{context}领域的简介
- 该品牌{context}产品的优势与劣势
- 与其他品牌{context}产品的对比（以表格呈现）
- 总结建议
- 该品牌的官方网站

在回答的最后，请按以下格式列出该品牌在{context}领域最具代表性的5个产品系列（如果不足5个则列出全部）：

◆◆◆ 产品系列1（仅限{context}相关）
◆◆◆ 产品系列2（仅限{context}相关）
◆◆◆ 产品系列3（仅限{context}相关）
◆◆◆ 产品系列4（仅限{context}相关）
◆◆◆ 产品系列5（仅限{context}相关）

重要提示：
1. 只列出与{context}相关的产品系列
2. 如果该品牌没有{context}相关产品，请明确说明
3. 不要列出其他类型的产品系列（如手机、电视等与{context}无关的产品）
4. 产品系列名称前必须使用"◆◆◆"标记，不要使用其他符号
"""
        if language != "zh":
            base_prompt += "\n5. Product series names must be in English, not Chinese"

        full_prompt = base_prompt.replace("请用中文回答。", "").strip() + lang_instruction

    stream = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "user", "content": full_prompt}
        ],
        stream=True
    )

    for chunk in stream:
        if chunk.choices[0].delta.content is not None:
            yield chunk.choices[0].delta.content


def call_deepseek_product_prompt(brand_product: str, language: str = "zh") -> str:
    from docx import Document
    try:
        document = Document("key/deepseek_product_prompt.docx")
        prompt_text = "\n".join([para.text for para in document.paragraphs])
    except:
        prompt_text = """你将被用于回答某个具体产品系列的详细信息。请根据输入的"品牌名 + 产品系列名"，输出以下内容：
- 产品系列简介
- 主要型号和配置
- 价格区间
- 核心特点和卖点
- 目标人群
- 与竞品的对比优势
- 购买建议"""

    # 添加语言指令
    lang_instruction = LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["zh"])
    full_prompt = prompt_text + f"\n\n品牌和产品系列：{brand_product}" + lang_instruction

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "user", "content": full_prompt}
        ]
    )
    return response.choices[0].message.content


def call_deepseek_product_prompt_stream(brand_product: str, language: str = "zh"):
    """流式调用产品系列详情API"""
    from docx import Document
    try:
        document = Document("key/deepseek_product_prompt.docx")
        prompt_text = "\n".join([para.text for para in document.paragraphs])
    except:
        prompt_text = """你将被用于回答某个具体产品系列的详细信息。请根据输入的"品牌名 + 产品系列名"，输出以下内容：
- 产品系列简介
- 主要型号和配置
- 价格区间
- 核心特点和卖点
- 目标人群
- 与竞品的对比优势
- 购买建议"""

    lang_instruction = LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["zh"])
    full_prompt = prompt_text + f"\n\n品牌和产品系列：{brand_product}" + lang_instruction

    stream = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "user", "content": full_prompt}
        ],
        stream=True
    )

    for chunk in stream:
        if chunk.choices[0].delta.content is not None:
            yield chunk.choices[0].delta.content


def call_deepseek_product_with_context_prompt(brand_product: str, context: str, language: str = "zh") -> str:
    """根据商品类型上下文生成产品系列详情"""
    # 添加语言指令
    lang_instruction = LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["zh"])

    full_prompt = f"""
你将被用于回答某个具体产品系列的详细信息。

品牌和产品系列：{brand_product}
商品类型背景：{context}

请根据输入的"品牌名 + 产品系列名"，输出以下内容：

- 产品系列简介（着重介绍与{context}相关的功能）
- 主要型号和配置
- 价格区间
- 核心特点和卖点
- 目标人群
- 与竞品的对比优势
- 购买建议

请用简洁明了的语言回答，重点突出该产品系列的特色和适用场景。
""" + lang_instruction

    response = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "user", "content": full_prompt}
        ]
    )
    return response.choices[0].message.content


def call_deepseek_product_with_context_prompt_stream(brand_product: str, context: str, language: str = "zh"):
    """流式调用带上下文的产品系列详情API"""
    lang_instruction = LANGUAGE_INSTRUCTIONS.get(language, LANGUAGE_INSTRUCTIONS["zh"])

    full_prompt = f"""
你将被用于回答某个具体产品系列的详细信息。

品牌和产品系列：{brand_product}
商品类型背景：{context}

请根据输入的"品牌名 + 产品系列名"，输出以下内容：

- 产品系列简介（着重介绍与{context}相关的功能）
- 主要型号和配置
- 价格区间
- 核心特点和卖点
- 目标人群
- 与竞品的对比优势
- 购买建议

请用简洁明了的语言回答，重点突出该产品系列的特色和适用场景。
""" + lang_instruction

    stream = client.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {"role": "user", "content": full_prompt}
        ],
        stream=True
    )

    for chunk in stream:
        if chunk.choices[0].delta.content is not None:
            yield chunk.choices[0].delta.content